from flask import Flask, request, render_template, redirect, url_for
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
from rank_bm25 import BM25Okapi
import os
import re
from werkzeug.utils import secure_filename
import csv
import PyPDF2
from docx import Document

app = Flask(__name__)

DOCUMENTS_DIR = 'documents/'
UPLOAD_FOLDER = 'static/images/'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

if not os.path.exists(DOCUMENTS_DIR):
    os.makedirs(DOCUMENTS_DIR)
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_documents():
    docs = {}
    for filename in os.listdir(DOCUMENTS_DIR):
        file_path = os.path.join(DOCUMENTS_DIR, filename)
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        if file_extension == "txt":
            with open(file_path, 'r', encoding='utf-8') as file:
                docs[filename] = file.read()
        elif file_extension == "pdf":
            docs[filename] = read_pdf(file_path)
        elif file_extension == "docx":
            docs[filename] = read_docx(file_path)
    
    return docs

def read_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text

def read_docx(file_path):
    doc = Document(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text
    return text


def load_stopwords():
    stopwords = set()
    with open('stopwordbahasa.csv', newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            stopwords.add(row[0].strip())
    return stopwords
STOPWORDS = load_stopwords()

@app.route('/upload_document', methods=['POST'])
def upload_document():
    title = request.form['title']
    content = request.form['content']

    if 'image' not in request.files:
        return "No image part"
    
    image_file = request.files['image']
    
    if image_file and allowed_file(image_file.filename):
        filename = secure_filename(image_file.filename)
        image_path = os.path.join(UPLOAD_FOLDER, f"{title}.jpg")
        image_file.save(image_path)

        text_file_path = os.path.join(DOCUMENTS_DIR, f"{title}.txt")
        with open(text_file_path, 'w', encoding='utf-8') as text_file:
            text_file.write(content)
        
        return redirect(url_for('index'))
    else:
        return "Invalid file type or no file uploaded"

def load_document_by_name(filename):
    file_path = os.path.join(DOCUMENTS_DIR, filename)
    file_extension = filename.split('.')[-1].lower()
    
    if file_extension == 'txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    elif file_extension == 'pdf':
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            return f"Error reading PDF: {e}"
    
    elif file_extension == 'docx':
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            return text
        except Exception as e:
            return f"Error reading DOCX: {e}"
    
    else:
        return f"Unsupported file type: {file_extension}"

    
def preprocess_text(text):
    text = text.lower()
    tokens = re.findall(r'\b\w+\b', text)
    filtered_tokens = []
    detected_stopwords = []
    
    for token in tokens:
        if token in STOPWORDS:
            detected_stopwords.append(token)
        else:
            filtered_tokens.append(token)
    
    factory = StemmerFactory()
    stemmer = factory.create_stemmer()
    stemmed_tokens = [stemmer.stem(word) for word in filtered_tokens]
    
    return stemmed_tokens, detected_stopwords

def get_bm25_similarity(query, documents):
    processed_query, _ = preprocess_text(query)
    
    processed_docs = [preprocess_text(doc_text)[0] for doc_text in documents.values()]
    bm25 = BM25Okapi(processed_docs)
    query_bm25_score = bm25.get_scores(processed_query)
    
    return sorted(zip(documents.keys(), query_bm25_score), key=lambda x: x[1], reverse=True)


@app.route('/')
def index():
    documents = load_documents()
    return render_template('index.html', documents=documents)

@app.route('/document/<filename>')
def document(filename):
    content = load_document_by_name(filename)
    processed_tokens, detected_stopwords = preprocess_text(content)
    filename_without_extension = os.path.splitext(filename)[0]
    word_count = len(processed_tokens)
    stopword_count = len(detected_stopwords)
    word_frequency = {word: processed_tokens.count(word) for word in set(processed_tokens)}
    stopword_frequency = {word: detected_stopwords.count(word) for word in set(detected_stopwords)}
    unique_word_count = len(set(processed_tokens))
    unique_stopword_count = len(set(detected_stopwords))
    
    return render_template('document.html', content=content, filename=filename_without_extension,
                           word_count=word_count, word_frequency=word_frequency,
                           unique_word_count=unique_word_count, stopword_frequency=stopword_frequency, stopword_count=stopword_count, unique_stopword_count=unique_stopword_count)


@app.route('/search', methods=['GET', 'POST'])
def search():
    documents = load_documents()
    if request.method == 'POST':
        query = request.form['query']
        similarity_scores = get_bm25_similarity(query, documents)
        return render_template('results.html', similarity_scores=similarity_scores, query=query)
    return render_template('index.html', documents=documents)

if __name__ == '__main__':
    app.run(debug=True)
